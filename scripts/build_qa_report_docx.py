from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "qa-business-feature-report.md"
OUTPUT = ROOT / "docs" / "qa-business-feature-report.docx"


COLORS = {
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "1F2937",
    "muted": "666666",
    "border": "D9E2EF",
    "header_fill": "F2F4F7",
    "callout_fill": "E8EEF5",
    "risk_fill": "FFF0F0",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EF", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quizzy FE QA Report")
    set_run_font(run, 9, False, COLORS["muted"])


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("QA Business & Feature Test Report")
    set_run_font(run, 22, True, COLORS["blue"])
    title.paragraph_format.space_after = Pt(3)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Quizzy FE - report for development team")
    set_run_font(run, 11, False, COLORS["muted"])
    subtitle.paragraph_format.space_after = Pt(12)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    set_table_width(meta)
    set_table_borders(meta)
    rows = [
        ("Ngay test", "07/07/2026"),
        ("Moi truong", "http://localhost:3000"),
        ("Backend", "https://quizzybe-production.up.railway.app/"),
        ("Tai khoan test", "student@gizmo.local"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].width = Inches(1.65)
        row.cells[1].width = Inches(4.85)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], COLORS["header_fill"])
        r0 = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(r0, 9.5, True, COLORS["dark_blue"])
        r1 = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r1, 9.5, False, COLORS["ink"])


def add_callout(doc: Document, text: str, fill: str = "E8EEF5") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="C7D4E7")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 10.5, True, COLORS["dark_blue"])
    doc.add_paragraph()


def clean_inline_markdown(text: str) -> str:
    text = text.replace("**", "")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def add_markdown_paragraph(doc: Document, text: str, style=None) -> None:
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`|\\*\\*[^*]+\\*\\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_run_font(run, 10, False, COLORS["dark_blue"])
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_run_font(run, 11, True, COLORS["ink"])
        else:
            run = p.add_run(part)
            set_run_font(run, 11, False, COLORS["ink"])


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [clean_inline_markdown(cell.strip()) for cell in raw.strip("|").split("|")]
        if not all(re.fullmatch(r"-+", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table)

    if col_count == 6:
        widths = [0.75, 1.25, 1.45, 1.45, 0.85, 0.75]
    elif col_count == 2:
        widths = [1.9, 4.6]
    else:
        widths = [6.5 / col_count] * col_count

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, COLORS["header_fill"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(
                run,
                8 if col_count >= 6 else 9.5,
                r_idx == 0,
                COLORS["dark_blue"] if r_idx == 0 else COLORS["ink"],
            )
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    set_table_borders(table, color="CBD5E1")
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(COLORS["dark_blue"])
    doc.add_paragraph()


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    # Skip the original H1 and metadata because the DOCX has a custom title block.
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("# QA Business"):
            i += 1
            continue
        if stripped.startswith("Ngay test:") or stripped.startswith("Moi truong:") or stripped.startswith("Backend dang") or stripped.startswith("Tai khoan test:") or stripped.startswith("Trang/feature"):
            i += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table_from_rows(doc, rows)
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline_markdown(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            title = clean_inline_markdown(stripped[3:])
            doc.add_heading(title, level=1)
            if title == "2. Executive Summary":
                add_callout(
                    doc,
                    "Key blockers: public deck access, My decks data display, and Learn/Test feedback accuracy.",
                    COLORS["callout_fill"],
                )
        elif stripped.startswith("# "):
            doc.add_heading(clean_inline_markdown(stripped[2:]), level=1)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            run = p.add_run(clean_inline_markdown(re.sub(r"^\d+\.\s+", "", stripped)))
            set_run_font(run, 11, False, COLORS["ink"])
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(clean_inline_markdown(stripped[2:]))
            set_run_font(run, 11, False, COLORS["ink"])
        else:
            add_markdown_paragraph(doc, stripped)
        i += 1


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_document(doc)
    add_footer(doc)
    add_title_block(doc)
    add_markdown_body(doc, markdown)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
